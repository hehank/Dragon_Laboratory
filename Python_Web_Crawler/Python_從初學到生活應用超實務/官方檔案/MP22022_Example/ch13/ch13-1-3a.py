import docx

doc = docx.Document("word文件.docx")
doc.add_heading("Python程式設計", level=1)
doc.add_heading("生活應用實務", level=2)
doc.save("word文件1.docx")
 


