import docx

doc = docx.Document("Python開發環境.docx")
print("段落數: ", len(doc.paragraphs))
for para in doc.paragraphs:
    print("連續文字數: ", len(para.runs))


