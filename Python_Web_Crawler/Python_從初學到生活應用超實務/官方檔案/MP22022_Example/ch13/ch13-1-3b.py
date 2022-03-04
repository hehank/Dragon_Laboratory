import docx

doc = docx.Document("word文件1.docx")
para = doc.paragraphs[1]
para.insert_paragraph_before("Python是一種直譯語言。")
para1 = doc.add_paragraph("Python程式是使用直譯器一行一行轉換成機器語言後，馬上執行程式碼。")
para1.add_run("WinPython")
para1.add_run(", ")
para1.add_run("Anaconda")
doc.save("word文件2.docx")
 


