import docx

doc = docx.Document("word文件6.docx")
print("段落數: ", len(doc.paragraphs))
para2 = doc.paragraphs[1]
para3 = doc.paragraphs[2]
run = para2.add_run("-第一版")
run.bold = True
para3.runs[0].italic = True
doc.save("word文件7.docx") 


